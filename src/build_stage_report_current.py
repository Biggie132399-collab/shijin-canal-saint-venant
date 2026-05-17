#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "阶段性汇报_0-456渠段SaintVenant与PatiRSR适用性分析_按老师意见修订.docx"

FIG_DIR = ROOT / "results" / "saint_venant_dispatch_results" / "figures"
PATI_DIR = ROOT / "results" / "pati_local_reach_applicability_results"
PATI_FIG_DIR = PATI_DIR / "figures"

FIGURES = {
    "fig1": FIG_DIR / "fig1_geographic_topology_main_canal_branches.png",
    "fig2": FIG_DIR / "fig2_diversion_outflow_process.png",
    "fig3": FIG_DIR / "fig3_cumulative_supply_process.png",
    "fig4": FIG_DIR / "fig4_key_node_depth_process.png",
    "fig5": FIG_DIR / "fig6_max_depth_envelope.png",
    "pati": PATI_FIG_DIR / "fig_pati_local_reach_applicability.png",
    "pati_table": PATI_FIG_DIR / "fig_pati_local_reach_applicability_table.png",
}

DISPATCH_CSV = ROOT / "results" / "saint_venant_dispatch_results" / "dispatch_feasibility_outlets.csv"
PATI_CSV = PATI_DIR / "pati_local_reach_applicability.csv"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D8E0EA")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_formula(doc: Document, number: int, expr: str, desc: str | None = None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [7600, 1100])
    expr_cell, no_cell = table.rows[0].cells
    expr_cell.text = ""
    p = expr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(expr)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r.font.size = Pt(11)
    no_cell.text = ""
    p2 = no_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"Eq. ({number})")
    r2.font.name = "等线"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r2.font.size = Pt(10)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if desc:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.5)
        p3.paragraph_format.right_indent = Cm(0.5)
        p3.paragraph_format.space_after = Pt(6)
        r3 = p3.add_run(desc)
        r3.font.name = "等线"
        r3._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(88, 101, 121)


def add_variable_lines(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        r = p.add_run(line)
        r.font.name = "等线"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        r.font.size = Pt(10)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(88, 101, 121)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(doc: Document, image_path: Path, caption: str, width_in: float = 6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    add_caption(doc, caption)


def read_dispatch_rows():
    rows = []
    with DISPATCH_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def read_pati_rows():
    rows = []
    with PATI_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def add_dispatch_table(doc: Document):
    rows = read_dispatch_rows()
    headers = ["分水口", "Qmax(m³/s)", "需水量(m³)", "起流(h)", "关闸(h)", "供水量(m³)", "是否满足"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [850, 1150, 1500, 1050, 1050, 1500, 1100]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8, color="263241")
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        values = [
            row["node"],
            f"{float(row['specified_qmax_m3s']):.0f}",
            f"{float(row['demand_m3']):.0f}",
            f"{float(row['first_positive_time_h']):.2f}",
            f"{float(row['close_time_h']):.2f}",
            f"{float(row['final_supplied_m3']):.0f}",
            "满足" if row["demand_satisfied"] == "True" else "未满足",
        ]
        for i, val in enumerate(values):
            set_cell_text(cells[i], val, size=8)
    set_table_widths(table, widths)
    set_table_borders(table)
    doc.add_paragraph()


def add_pati_table(doc: Document):
    rows = read_pati_rows()
    headers = ["反演段", "长度(km)", "h RMSE(m)", "h NSE", "Q RMSE(m³/s)", "Q NSE", "结论"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1200, 1050, 1200, 1100, 1400, 1100, 1500]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        values = [
            row["reach"],
            f"{float(row['length_m'])/1000:.2f}",
            f"{float(row['depth_rmse_m']):.3f}",
            f"{float(row['depth_nse']):.3f}",
            f"{float(row['q_rmse_m3s']):.2f}",
            f"{float(row['q_nse']):.3f}",
            row["applicability"],
        ]
        for i, val in enumerate(values):
            color = "2A9D8F" if (i == 6 and val == "仅水深可参考") else None
            set_cell_text(cells[i], val, size=8, color=color)
    set_table_widths(table, widths)
    set_table_borders(table)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [9000])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title + "：")
    r.bold = True
    r.font.name = "等线"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 77, 120)
    r2 = p.add_run(body)
    r2.font.name = "等线"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r2.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)
    doc.add_paragraph()


def setup_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "等线"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 8),
        ("Subtitle", 11, "586579", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "等线"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "等线"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc: Document):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("阶段性汇报 | 0–456渠段水动力模拟与Pati RSR适用性分析")
    r.font.name = "等线"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(88, 101, 121)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("0–456主干渠动态供配水模拟与 Pati RSR 适用性检验阶段性汇报")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"按老师意见补充数值离散、分水源项、公式编号与参考文献 | {date.today().isoformat()}")

    add_callout(
        doc,
        "阶段性结论",
        "当前结果支持将一维 Saint-Venant 正演作为0–456渠段动态调度主模型；Pati RSR 在简单规则渠段可闭合，在少数局部渠段对水深反演有参考价值，但流量反演和全渠段长距离反推不稳定，因此不宜作为复杂渠系主模型。",
        "E8EEF5",
    )

    add_heading(doc, "1. 研究对象与当前建模目标", 1)
    add_body(
        doc,
        "本阶段聚焦石津灌区前约20 km的0–456主干渠段，考虑渠首来水、多个配水口取水、关键节点水深变化与渠段安全水位约束。当前目标不是直接形成最终调度优化模型，而是先建立可解释的物理基准：用一维 Saint-Venant 正演描述渠首来水波沿主渠传播、配水口出流和沿程水位动态，并检验 Pati et al. (2023) reverse-stage routing 方法在本渠段条件下的适用边界。"
    )
    add_bullet(doc, "渠首边界：流量由0平滑爬升至80 m³/s。")
    add_bullet(doc, "配水口能力：71、89口为20 m³/s，287口为12 m³/s，150、194、349、383口为5 m³/s。")
    add_bullet(doc, "配水口控制：按最大能力放水，累计供水达到假设需水量后关闸。")
    add_bullet(doc, "水位约束：关键节点和沿程最大水深与渠深/安全水深进行比较。")

    add_heading(doc, "2. Saint-Venant 正演模型与数值离散", 1)
    add_body(
        doc,
        "本次报告将原先 Muskingum-Cunge 快速验证思路调整为一维 Saint-Venant 动态正演框架。当前计算程序为本人在 Python 中编写的原型程序，未直接调用开源水动力求解器；主要代码文件列于下方。"
    )
    add_bullet(doc, "/Users/pupu/Documents/New project/stage7_saint_venant_fig2_revised.py：Saint-Venant 正演与配水口出流过程。")
    add_bullet(doc, "/Users/pupu/Documents/New project/stage6_saint_venant_fig3.py：累计供水量积分与图3绘制。")
    add_bullet(doc, "/Users/pupu/Documents/New project/stage6_saint_venant_fig4_key_depth.py：关键节点水深动态过程。")
    add_bullet(doc, "/Users/pupu/Documents/New project/saint_venant_dispatch_fig6.py：沿程最大水深包络线。")
    add_body(
        doc,
        "为回应“公式之间需要相互衔接”的问题，本节按守恒变量、控制方程、数值通量、地形梯度项、摩阻项、分水源项和边界条件的顺序给出完整计算链条。"
    )
    add_heading(doc, "2.1 控制方程与守恒变量", 2)
    add_formula(doc, 1, "U = [A, Q]^T")
    add_formula(doc, 2, "∂U/∂t + ∂F(U)/∂x = S_b + S_f + S_div")
    add_formula(doc, 3, "F(U) = [Q, Q²/A + gI₁]^T")
    add_formula(doc, 4, "I₁(h) = (1/2)bh² + (1/3)zh³")
    add_variable_lines(
        doc,
        [
            "U：守恒变量向量；A：过水面积，单位 m²；Q：流量，单位 m³/s；F(U)：通量向量。",
            "t：时间，单位 s；x：沿程距离，单位 m；g：重力加速度，单位 m/s²。",
            "I₁：断面静水压力积分项，单位 m³；b：底宽，单位 m；z：边坡系数，表示水平:垂直；h：水深，单位 m。",
            "S_b：地形/床坡源项；S_f：Manning 摩阻源项；S_div：分水口侧向出流源项。",
        ],
    )

    add_heading(doc, "2.2 断面几何关系", 2)
    add_body(doc, "根据 input.txt 和 lineParam.txt 中的渠深、底宽、边坡角和 Manning 糙率参数，主渠每个节点均按所属线段参数计算梯形断面水力要素。")
    add_formula(doc, 5, "A = bh + zh²")
    add_formula(doc, 6, "B = b + 2zh")
    add_formula(doc, 7, "P = b + 2h√(1+z²)")
    add_formula(doc, 8, "R = A/P")
    add_variable_lines(
        doc,
        [
            "B：自由水面宽，单位 m；P：湿周，单位 m；R：水力半径，单位 m。",
            "D：渠深，单位 m；安全水深以渠深或设定安全比例作为参照，用于图4和图5的水位超限检查。",
        ],
    )

    add_heading(doc, "2.3 有限体积离散与 HLL 数值通量", 2)
    add_body(doc, "对第 i 个控制体，采用有限体积格式进行时间推进。当前程序中通量项显式计算，摩阻项随后半隐式修正。")
    add_formula(doc, 9, "U_i^(n+1,*) = U_i^n − (Δt/Δx_i)(F_(i+1/2)^n − F_(i−1/2)^n) + Δt S_(b,i)^n + Δt S_(div,i)^n")
    add_formula(doc, 10, "F_HLL = [s_R F(U_L) − s_L F(U_R) + s_Ls_R(U_R − U_L)]/(s_R − s_L)")
    add_formula(doc, 11, "s_L = min(u_L − c_L, u_R − c_R)")
    add_formula(doc, 12, "s_R = max(u_L + c_L, u_R + c_R)")
    add_formula(doc, 13, "u = Q/A")
    add_formula(doc, 14, "c = √(gA/B)")
    add_formula(doc, 15, "CFL = max_i[(|u_i|+c_i)Δt/Δx_i] ≤ CFL_max")
    add_variable_lines(
        doc,
        [
            "Δt：时间步长，单位 s；Δx_i：第 i 个控制体长度，单位 m；上标 n 表示第 n 个时间层。",
            "U_L、U_R：界面左右两侧状态量；F_(i+1/2)：第 i 与 i+1 个控制体之间的界面数值通量。",
            "u：断面平均流速，单位 m/s；c：明渠浅水波速，单位 m/s。这里采用 c=√(gA/B)，即以水力深度 A/B 代替矩形渠近似中的 h。",
            "显式通量项的稳定性检查采用 |u|+c，而不是仅按实际流速或仅按重力波速估算。当前绘图程序内部计算时间步长为1 s，输出间隔为60 s。",
        ],
    )

    add_heading(doc, "2.4 地形梯度项离散", 2)
    add_body(
        doc,
        "老师特别指出需要说明地形梯度项的离散。当前程序将渠底高程 Z_b 显式进入床坡源项，采用相邻节点渠底高程差计算局部床坡。内点采用中心差分，边界点采用单边差分。"
    )
    add_formula(doc, 16, "S_(b,i) = [0, gA_iS_(0,i)]^T")
    add_formula(doc, 17, "S_(0,i) = (Z_(b,i−1) − Z_(b,i+1))/(x_(i+1) − x_(i−1)),  1 ≤ i ≤ N−1")
    add_formula(doc, 18, "S_(0,0) = (Z_(b,0) − Z_(b,1))/(x_1 − x_0)")
    add_formula(doc, 19, "S_(0,N) = (Z_(b,N−1) − Z_(b,N))/(x_N − x_(N−1))")
    add_variable_lines(
        doc,
        [
            "Z_b：渠底高程，单位 m；S_0：局部床坡，无量纲。",
            "该处理使床坡由真实节点高程进入 Saint-Venant 动量方程，而不是使用单一常数坡降。",
            "需要说明的是，当前版本属于局部床坡源项离散，尚未实现严格 well-balanced 的 hydrostatic reconstruction；后续若用于更高精度水面线保持，可进一步升级。",
        ],
    )

    add_heading(doc, "2.5 Manning 摩阻项与半隐式处理", 2)
    add_body(doc, "Manning 摩阻坡降按下式计算，并在动量方程中作为耗散源项。为降低显式摩阻带来的振荡，当前程序对摩阻项采用半隐式修正。")
    add_formula(doc, 20, "S_f = n²Q|Q|/(A²R^(4/3))")
    add_formula(doc, 21, "S_f^Q = [0, −gAS_f]^T")
    add_formula(doc, 22, "Q_i^(n+1) = Q_i^(n+1,*)/[1 + Δt·g·n²|Q_i^n|/(A_i^(n+1,*)R_i^(4/3))]")
    add_variable_lines(
        doc,
        [
            "n：Manning 糙率系数，单位 s/m^(1/3)；S_f：摩阻坡降，无量纲。",
            "Q_i^(n+1,*)：未进行摩阻修正的中间流量；Q_i^(n+1)：半隐式摩阻修正后的流量。",
        ],
    )

    add_heading(doc, "2.6 分水口在 Saint-Venant 框架下的描述", 2)
    add_body(
        doc,
        "分水口不是单独脱离 Saint-Venant 方程的经验后处理，而是作为集中侧向出流源项进入连续方程和动量方程。对位于节点 i 的第 k 个分水口，其实际出流从主渠控制体中扣除，并同步扣除相应动量。"
    )
    add_formula(doc, 23, "S_(div,i) = [−Q_(div,k)/Δx_i, −βQ_(div,k)u_i/Δx_i]^T")
    add_formula(doc, 24, "Q_(div,k)(t) = min[φ(h_i)min(Q_(max,k),Q_(safe,k)), Q_(stor,k)(t), W_(rem,k)(t)/Δt]")
    add_formula(doc, 25, "Q_(safe,k) = (1/n_b)A_b(0.9D_b)R_b(0.9D_b)^(2/3)S_(b,k)^(1/2)")
    add_formula(doc, 26, "φ(h_i)=0, h_i≤0.2D_i;  φ(h_i)=√[(h_i−0.2D_i)/(0.6D_i−0.2D_i)], 0.2D_i<h_i<0.6D_i;  φ(h_i)=1, h_i≥0.6D_i")
    add_formula(doc, 27, "W_k(t_m) = Σ_(j=1..m) Q_(div,k)(t_j)Δt")
    add_variable_lines(
        doc,
        [
            "Q_div,k：第 k 个配水口实际出流，单位 m³/s；Q_max,k：第 k 个配水口设定最大过流能力，单位 m³/s。",
            "Q_safe,k：支渠安全过流能力估计值，单位 m³/s；D_b、A_b、R_b、n_b 分别为支渠首段渠深、面积、水力半径和糙率。",
            "Q_stor,k：该主渠控制体在最小湿润水深以上可释放的水量折算流量，单位 m³/s。",
            "W_rem,k：剩余需水量，单位 m³；W_k：累计供水量，单位 m³。",
            "φ(h_i)：水深启闭系数，无量纲；D_i：节点 i 所属渠段渠深，单位 m；β：侧向出流动量修正系数。当前代码取 β=1，即按本地平均流速扣除主流动量；若后续取得闸门角度和支渠流向资料，可改为校准参数。",
        ],
    )

    add_heading(doc, "2.7 边界条件与调度工况", 2)
    add_formula(doc, 28, "Q_0(t) = Q_in(t)")
    add_formula(doc, 29, "A_0(t) = A[h_n(Q_in,S_0,n,b,z)]")
    add_formula(doc, 30, "Q_(N+1)(t) = Q_N(t)")
    add_variable_lines(
        doc,
        [
            "渠首上游边界采用给定流量过程 Q_in(t)，本工况由0平滑爬升至80 m³/s。",
            "渠首 ghost cell 的面积由 Manning 正常水深 h_n 估算，用于构造入口状态。",
            "下游边界采用零梯度出流近似，即末端外侧 ghost cell 流量取末端控制体流量。",
        ],
    )

    add_heading(doc, "3. 主要结果图", 1)
    add_figure(doc, FIGURES["fig1"], "图1  0–456主渠段和配水口位置图。该图展示主渠节点空间走向及71、89、150、194、287、349、383等配水口位置。")
    add_figure(doc, FIGURES["fig2"], "图2  各配水口实际放水过程。曲线表示各口实际出流，虚线表示达到需水量后的关闸时刻。")
    add_body(doc, "配水口供水完成情况如下表所示。所有假设需水量均在本次工况中得到满足。")
    add_dispatch_table(doc)
    add_figure(doc, FIGURES["fig3"], "图3  各配水口累计供水量与需水量对比图。曲线变平表示该配水口累计供水达到目标并执行关闸。")
    add_figure(doc, FIGURES["fig4"], "图4  关键节点水深动态变化图。图中同时给出安全水深/渠深参照，用于判断动态过程中是否出现水位超限。")
    add_figure(doc, FIGURES["fig5"], "图5  沿程最大水深包络线。结果用于检查0–456主渠段全程最大水深是否超过渠段限制水深。")

    add_heading(doc, "4. Pati RSR 方法复核与局部适用性评价", 1)
    add_body(
        doc,
        "Pati et al. (2023) 的 RSR 方法以 stage（水位）作为反演主变量。原文先利用 HEC-RAS 在规则矩形/梯形河道中进行正向演算，以正演得到的下游水位过程作为反向演算边界，再由 RSR 模型反推上游水位和流量。该方法在规则或等效规则河段中具有较好表现，但其适用性依赖下游 stage 边界一致性、等效断面参数、Manning 糙率、侧向流处理和适用性判据。"
    )
    add_body(
        doc,
        "本课题对象是包含多个分水口和闸门控制的灌区干渠。与天然河道两水文站之间的 reach 相比，本渠段存在集中侧向出流、调度波、关闸扰动和安全水位约束。因此，本阶段不将 Pati RSR 作为全渠段主模型，而将其定位为局部 stage 反演与适用性诊断工具。"
    )
    add_body(doc, "Pati RSR 的核心面积梯度关系为：")
    add_formula(doc, 31, "∂A/∂x ≈ (1/c²)∂Q/∂t − (2/c)∂A/∂t + q_l/c")
    add_body(doc, "在单个局部等效渠段内，由下游断面向上游断面反推时，可写成如下离散形式：")
    add_formula(doc, 32, "A_u,j = A_d,j − (Δx/c_j²)(∂Q_d/∂t)_j + (2Δx/c_j)(∂A_d/∂t)_j − Δx q_l/c_j")
    add_body(doc, "波速 c 采用 Pati et al. (2023) 的扩散波波速表达：")
    add_formula(doc, 33, "c = (Q/A)[5/3 − (2/3)(R/B)(dP/dh)]")
    add_variable_lines(
        doc,
        [
            "A_u,j：第 j 个时间层的上游过水面积，单位 m²；A_d,j：下游过水面积，单位 m²。",
            "Δx：局部反演渠段长度，单位 m；c_j：第 j 个时间层扩散波波速，单位 m/s。",
            "Q_d：下游断面流量，单位 m³/s；q_l：单位长度侧向流量，单位 m²/s；本次分水口间局部渠段内部无分水口，故取 q_l=0。",
            "dP/dh：湿周对水深的导数，无量纲；对于梯形断面，dP/dh=2√(1+z²)。",
        ],
    )
    add_body(doc, "反演结果采用 RMSE 与 NSE 进行评价：")
    add_formula(doc, 34, "RMSE = √[(1/N) Σᵢ(y_sim,i − y_obs,i)²]")
    add_formula(doc, 35, "NSE = 1 − Σᵢ(y_sim,i − y_obs,i)² / Σᵢ(y_obs,i − ȳ_obs)²")
    add_variable_lines(
        doc,
        [
            "y_sim,i：第 i 个时间点的反演值；y_obs,i：Saint-Venant 正演基准值；ȳ_obs：基准序列均值；N：样本数量。",
            "RMSE 越接近 0 表示误差越小；NSE=1 表示完全一致；NSE=0 表示与均值基准相当；NSE<0 表示反演效果劣于均值基准。",
        ],
    )
    add_figure(doc, FIGURES["pati"], "图6  Pati RSR 局部适用性分段评价图。各分水口之间构造局部等效梯形渠段，用下游stage反推上游stage，并与Saint-Venant正演真值比较。")
    add_body(doc, "局部适用性评价表明，只有89→71和349→287两段达到“仅水深可参考”；没有任何分段同时达到水深和流量均可接受。")
    add_pati_table(doc)
    add_callout(
        doc,
        "Pati RSR 结论",
        "Pati RSR 在本渠系中不宜表述为全渠段成功反演模型。更稳妥的表述是：其在规则或局部等效渠段中可用于stage反演尝试，但在复杂灌区干渠长距离反推与流量反演中存在明显误差放大。",
        "FFF8E6",
    )

    add_heading(doc, "5. 阶段性判断", 1)
    add_bullet(doc, "Saint-Venant 正演模型已经能够生成配水口出流、累计供水、关键节点水深和沿程最大水深包络线，可作为当前主物理模型。")
    add_bullet(doc, "当前调度工况下，各配水口假设需水量均可满足，关键节点和沿程最大水深未显示出明显超限风险。")
    add_bullet(doc, "Pati RSR 在单一规则渠段上可以闭合，但在本渠系分段评价中仅少数渠段的水深反演可参考，流量反演整体不稳定。")
    add_bullet(doc, "因此，Pati RSR 更适合作为局部方法对比和物理诊断，不适合作为0–456复杂渠系调度主模型。")

    add_heading(doc, "6. 下一步思路", 1)
    add_body(
        doc,
        "下一阶段建议围绕“Saint-Venant 物理正演样本库 + 物理约束图神经网络代理模型”推进。Saint-Venant 模型继续作为物理一致的正演基准，用于生成不同渠首边界、需水量组合、分水口能力和闸门策略下的水位/流量响应样本；随后构建面向复杂渠系拓扑的 PI-GNN 代理响应模型，实现快速预测和调度方案筛选。"
    )
    add_bullet(doc, "完善 Saint-Venant 正演：加入更真实的闸门出流公式、分水口控制逻辑和安全水位约束。")
    add_bullet(doc, "生成样本库：设计多组渠首流量、各分水口需水量、不同开关闸策略的正演工况。")
    add_bullet(doc, "构建图结构：节点表示渠首、干渠节点、配水口和控制点；边表示相邻渠段，边特征包含长度、床坡、糙率和断面参数。")
    add_bullet(doc, "设计物理约束损失：在数据误差之外加入质量守恒、水位不超限、分水口供需约束等损失项。")
    add_bullet(doc, "形成毫秒级代理响应：用 PI-GNN 替代重复正演，用于快速评估多种调度方案。")

    add_heading(doc, "7. 推荐汇报表述", 1)
    add_body(
        doc,
        "当前成果可以概括为：本文已建立0–456主干渠一维 Saint-Venant 动态供配水正演模型，能够模拟渠首来水波传播、配水口实际出流、累计供水和关键节点水深变化。Pati RSR 方法经复核后，在规则或局部等效渠段中对水深反演具有一定参考价值，但在本渠系全渠段和流量反演任务中存在适用性限制。因此后续以 Saint-Venant 正演作为物理基准，进一步构建融合渠系拓扑和水动力约束的 PI-GNN 代理响应模型。"
    )

    add_heading(doc, "8. 参考文献与代码说明", 1)
    add_body(doc, "本报告未直接调用 HEC-RAS、SWMM、TELEMAC、Clawpack 等开源或商业水动力求解器；当前图件由自编 Python 原型程序生成。后续若接入开源代码或商业软件，应在正文中明确软件名称、版本、求解格式和调用方式。")
    add_bullet(doc, "Saint-Venant, A. J. C. B. (1871). Théorie du mouvement non permanent des eaux. Comptes rendus de l'Académie des Sciences.")
    add_bullet(doc, "Chow, V. T. (1959). Open-Channel Hydraulics. McGraw-Hill.")
    add_bullet(doc, "Harten, A., Lax, P. D., & van Leer, B. (1983). On upstream differencing and Godunov-type schemes for hyperbolic conservation laws. SIAM Review, 25(1), 35–61.")
    add_bullet(doc, "U.S. Army Corps of Engineers. (2016). HEC-RAS River Analysis System Hydraulic Reference Manual, Version 5.0.")
    add_bullet(doc, "Pati, A., et al. (2023). A physically-based reverse-stage routing model considering lateral flow. Water Resources Research, 59, e2022WR034150.")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
